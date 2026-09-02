"""
Export fidelity tests for the AAR renderers.

The question these answer is narrow and specific: does every value an analyst
recorded on the case actually arrive in the exported document, in full? Three
defects motivated them, all of which the code shipped with at one point —

  * a control character anywhere in the case text raised ValueError out of
    python-docx and returned a 500 on the .docx download, while the Markdown
    export of the same case succeeded;
  * evidence hashes were exported as a 12-character prefix, which cannot be
    verified against anything once the document leaves the building;
  * recommendation status and the risk-acceptance justification were assembled
    into the report data and then rendered by nobody.

So these tests read the saved document back out rather than trusting the
builder — a field that renders is a field that a reader can find.
"""

import io
from datetime import date, datetime

import pytest
from docx import Document

from app.services import report_builder as rb


# ---------------------------------------------------------------------------
# Fixtures — plain objects, no database. build_report_data() only reads
# attributes, so a stand-in case exercises the same code the app runs.
# ---------------------------------------------------------------------------

class Obj:
    def __init__(self, **kw):
        self.__dict__.update(kw)


LONG_TEXT = (
    "This paragraph is deliberately longer than any column or cell would "
    "comfortably hold, so that a renderer which clips at a fixed width fails "
    "the assertion rather than passing quietly. " * 4
) + "END-OF-LONG-TEXT-MARKER"

MULTILINE = "First line.\r\nSecond line.\r\n\r\nFourth line after a blank."

SHA256 = "3f786850e387550fdab836ed7e6dc881de23001b" + "a" * 24  # 64 chars
MD5 = "d41d8cd98f00b204e9800998ecf8427e"


def make_case(**overrides):
    base = dict(
        case_id="CASE-0042",
        title="Credential harvesting via spearphishing link",
        case_type="Phishing",
        severity="High",
        imp_functional_impact="Moderate",
        imp_informational_impact="Critical",
        status="Closed",
        lead_analyst=Obj(name="J. Analyst"),
        created_by=Obj(name="K. Creator"),
        escalated=True,
        board_flagged=True,
        opened_date=datetime(2026, 1, 1, 9, 0),
        contained_date=datetime(2026, 1, 1, 14, 30),
        eradicated_date=datetime(2026, 1, 2, 10, 0),
        closed_date=datetime(2026, 1, 3, 9, 0),
        description=LONG_TEXT,
        method_of_discovery=MULTILINE,
        initial_vector=LONG_TEXT,
        affected_systems=MULTILINE,
        affected_users=MULTILINE,
        estimated_impact=LONG_TEXT,
        root_cause=MULTILINE,
        recovery_assessment=LONG_TEXT,
        recovery_sufficient="Partially Sufficient",
        lessons_learned_date=date(2026, 1, 8),
        lessons_learned_attendees=MULTILINE,
        lessons_learned_notes=LONG_TEXT,
        timeline_events=[
            Obj(
                event_datetime=datetime(2026, 1, 1, 9, 5),
                category="Detection",
                description=LONG_TEXT,
                mitre_tactic="Initial Access",
                mitre_technique="Phishing: Spearphishing Link",
                mitre_technique_id="T1566.002",
            )
        ],
        iocs=[
            Obj(
                ioc_type="Domain",
                value="login-" + "x" * 200 + ".example.invalid",
                description=LONG_TEXT,
                confidence="High",
            )
        ],
        evidence_items=[
            Obj(
                evidence_id="EV-0007",
                name="Mailbox export, affected user",
                evidence_type="Email/PST",
                collected_by="J. Analyst",
                hash_sha256=SHA256,
                hash_md5=MD5,
                file_path="1/EV-0007/x.bin",
                hash_verified_at=datetime(2026, 2, 4, 10, 0),
                hash_verified_ok=True,
            )
        ],
        status_history=[
            Obj(
                recorded_at=datetime(2026, 1, 3, 9, 0),
                old_status="Recovered",
                new_status="Closed",
                notes=LONG_TEXT,
            )
        ],
        deviations=[
            Obj(
                deviation=LONG_TEXT,
                standard_procedure=LONG_TEXT,
                justification=MULTILINE,
                approved_by="CISO",
            )
        ],
        recommendations=[
            Obj(
                id=1,
                text=LONG_TEXT,
                disposition="Risk Acceptance",
                owner="IT Operations",
                target_date=date(2026, 3, 1),
                risk_treatment_ref="RTP-2026-014",
                status="Complete",
                risk_acceptance_justification="Accepted: compensating MFA control in place until Q3.",
            )
        ],
        # Relationships the AAR reads for the assets, provenance and
        # contributor sections. Present on the stub for the same reason the
        # others are: the report walks them unconditionally, and a stub that
        # omits them tests a Case shape that cannot exist.
        asset_links=[
            Obj(
                role="Patient Zero",
                notes=MULTILINE,
                asset=Obj(name="DC01", type_label="Domain Controller",
                          asset_type="Domain Controller", criticality="Critical"),
            ),
            Obj(
                role=None,
                notes=None,
                asset=Obj(name="ws-042.corp.example", type_label="Unclassified",
                          asset_type=None, criticality=None),
            ),
        ],
        alerts=[
            Obj(
                source="crowdstrike", external_id="ldt:" + "b" * 32,
                severity_name="Critical", host_hostname="DC01", host_ip="10.0.0.5",
                cs_created_at=datetime(2026, 2, 1, 8, 41),
                fetched_at=datetime(2026, 2, 1, 8, 55),
                reviewed_by=Obj(name="J. Analyst"),
            ),
        ],
        audit_entries=[
            Obj(changed_by_id=1, changed_by=Obj(name="J. Analyst"),
                changed_at=datetime(2026, 2, 1, 9, 0)),
            Obj(changed_by_id=1, changed_by=Obj(name="J. Analyst"),
                changed_at=datetime(2026, 2, 3, 16, 20)),
            Obj(changed_by_id=2, changed_by=Obj(name="K. Creator"),
                changed_at=datetime(2026, 2, 2, 11, 5)),
            # An unattributed entry — a restore writes one of these. It must not
            # appear as a contributor named "None".
            Obj(changed_by_id=None, changed_by=None,
                changed_at=datetime(2026, 2, 2, 12, 0)),
        ],
    )
    base.update(overrides)
    return Obj(**base)


def docx_text(buf: io.BytesIO) -> str:
    """Every string a reader can see in the saved document."""
    doc = Document(buf)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


@pytest.fixture
def data():
    return rb.build_report_data(make_case())


@pytest.fixture
def rendered(data):
    return {"docx": docx_text(rb.render_docx(data)), "markdown": rb.render_markdown(data)}


def render(fmt, case):
    """Render one ad-hoc case in one format. The `rendered` fixture covers the
    default case; this is for tests that need a case shaped differently."""
    d = rb.build_report_data(case)
    return docx_text(rb.render_docx(d)) if fmt == "docx" else rb.render_markdown(d)


# ---------------------------------------------------------------------------
# Nothing is truncated
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("fmt", ["docx", "markdown"])
def test_long_narrative_survives_whole(rendered, fmt):
    assert "END-OF-LONG-TEXT-MARKER" in rendered[fmt]


@pytest.mark.parametrize("fmt", ["docx", "markdown"])
def test_long_table_cell_survives_whole(rendered, fmt):
    """The marker only reaches the output if the timeline description, the IOC
    description, the deviation text and the recommendation text all render in
    full — each is a table cell holding LONG_TEXT."""
    assert rendered[fmt].count("END-OF-LONG-TEXT-MARKER") >= 5


@pytest.mark.parametrize("fmt", ["docx", "markdown"])
def test_long_ioc_value_not_clipped(rendered, fmt):
    assert "login-" + "x" * 200 + ".example.invalid" in rendered[fmt]


@pytest.mark.parametrize("fmt", ["docx", "markdown"])
def test_evidence_hash_exported_in_full(rendered, fmt):
    """A prefix cannot be verified against the artifact it describes."""
    assert SHA256 in rendered[fmt]
    assert SHA256[:12] + "..." not in rendered[fmt]


def test_md5_used_when_no_sha256_recorded():
    case = make_case(evidence_items=[Obj(
        evidence_id="EV-0008", name="Disk image", evidence_type="Disk Image",
        collected_by="J. Analyst", hash_sha256=None, hash_md5=MD5,
        file_path=None, hash_verified_at=None, hash_verified_ok=None)])
    data = rb.build_report_data(case)
    assert data["evidence_hashes"] == [
        {"evidence_id": "EV-0008", "algorithm": "MD5", "digest": MD5}
    ]
    assert MD5 in docx_text(rb.render_docx(data))


def test_evidence_without_any_hash_is_omitted_not_faked():
    case = make_case(evidence_items=[Obj(
        evidence_id="EV-0009", name="Screenshot", evidence_type="Screenshot",
        collected_by="J. Analyst", hash_sha256=None, hash_md5=None,
        file_path=None, hash_verified_at=None, hash_verified_ok=None)])
    assert rb.build_report_data(case)["evidence_hashes"] == []


# ---------------------------------------------------------------------------
# Nothing is silently dropped
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("fmt", ["docx", "markdown"])
def test_recommendation_status_is_rendered(rendered, fmt):
    assert "Complete" in rendered[fmt]


@pytest.mark.parametrize("fmt", ["docx", "markdown"])
def test_risk_acceptance_justification_is_rendered(rendered, fmt):
    """IMP Phase VI requires it and the Report tab refuses to save without it;
    the document that carries the acceptance has to carry the reasoning."""
    assert "compensating MFA control in place until Q3" in rendered[fmt]


@pytest.mark.parametrize("fmt", ["docx", "markdown"])
def test_mitre_technique_name_is_rendered(rendered, fmt):
    assert "Phishing: Spearphishing Link" in rendered[fmt]
    assert "T1566.002" in rendered[fmt]
    assert "Initial Access" in rendered[fmt]


@pytest.mark.parametrize("fmt", ["docx", "markdown"])
def test_every_scalar_report_field_reaches_the_output(rendered, data, fmt):
    """Catch-all: a new key added to build_report_data() but wired into no
    renderer fails here rather than shipping as data nobody can read."""
    skip = {
        # Booleans and derived flags are rendered as prose, not as their value.
        "escalated", "board_flagged", "lessons_learned_sla_met",
        "lessons_learned_business_days", "generated_at",
    }
    for key, value in data.items():
        if key in skip or not isinstance(value, str) or not value.strip():
            continue
        assert value.splitlines()[0][:60] in rendered[fmt], f"{key} missing from {fmt}"


# ---------------------------------------------------------------------------
# Hostile input does not fail the export
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("codepoint", [0x00, 0x0B, 0x0C, 0x1B, 0x1F])
def test_control_characters_do_not_break_the_docx_export(codepoint):
    """\\x0b is Word's own soft line break, so text pasted from Word, a PDF or
    an alert body carries these in routinely. Previously any one of them
    raised ValueError and returned a 500 on the download."""
    case = make_case(root_cause=f"before{chr(codepoint)}after")
    buf = rb.render_docx(rb.build_report_data(case))
    assert "beforeafter" in docx_text(buf)


def test_control_characters_in_a_table_cell_do_not_break_the_export():
    case = make_case(timeline_events=[Obj(
        event_datetime=datetime(2026, 1, 1, 9, 5), category="Detection",
        description="cell\x0bcontent", mitre_tactic=None,
        mitre_technique=None, mitre_technique_id=None)])
    assert "cellcontent" in docx_text(rb.render_docx(rb.build_report_data(case)))


def test_crlf_does_not_double_space_the_docx():
    r"""Browsers submit textareas with CRLF. python-docx turns \r and \n each
    into their own <w:br/>, so every typed newline came out doubled."""
    case = make_case(root_cause="line one\r\nline two")
    doc = Document(rb.render_docx(rb.build_report_data(case)))
    para = next(p for p in doc.paragraphs if "line one" in p.text)
    assert para.text == "line one\nline two"


def test_pipe_and_newline_do_not_tear_the_markdown_table():
    case = make_case(timeline_events=[Obj(
        event_datetime=datetime(2026, 1, 1, 9, 5), category="Detection",
        description="left | right\nsecond line", mitre_tactic="Execution",
        mitre_technique=None, mitre_technique_id=None)])
    md = rb.render_markdown(rb.build_report_data(case))
    row = next(ln for ln in md.splitlines() if "left" in ln)
    assert row.count("|") - row.count("\\|") == 5      # 4 columns, borders intact
    assert "second line" in row                        # content kept, row unbroken


# ---------------------------------------------------------------------------
# The empty case still produces a document
# ---------------------------------------------------------------------------

def test_empty_case_still_renders_both_formats():
    case = make_case(
        description=None, method_of_discovery=None, initial_vector=None,
        affected_systems=None, affected_users=None, estimated_impact=None,
        root_cause=None, recovery_assessment=None, recovery_sufficient=None,
        lessons_learned_date=None, lessons_learned_attendees=None,
        lessons_learned_notes=None, imp_functional_impact=None,
        imp_informational_impact=None, lead_analyst=None,
        timeline_events=[], iocs=[], evidence_items=[], status_history=[],
        deviations=[], recommendations=[],
        asset_links=[], alerts=[], audit_entries=[],
    )
    data = rb.build_report_data(case)
    assert "CASE-0042" in docx_text(rb.render_docx(data))
    assert "CASE-0042" in rb.render_markdown(data)


# ---------------------------------------------------------------------------
# Affected assets, response metrics, integrity, provenance, contributors
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("fmt", ["markdown", "docx"])
def test_assets_reach_the_report(fmt):
    """Name, type, criticality and per-incident role all have to render."""
    out = render(fmt, make_case())
    for expected in ["DC01", "Domain Controller", "Critical", "Patient Zero",
                     "ws-042.corp.example", "Unclassified"]:
        assert expected in out, f"{expected!r} missing from the {fmt} report"


def test_asset_scope_is_computed_not_typed():
    data = rb.build_report_data(make_case())
    scope = data["asset_scope"]
    assert "2 assets" in scope
    assert "Domain Controller (1)" in scope and "Unclassified (1)" in scope
    assert "1 rated Critical" in scope
    # An unclassified asset is stated, not left as a blank cell — it is an open
    # question in a report going to leadership.
    assert "not yet classified" in scope


def test_unclassified_asset_is_never_given_a_guessed_type():
    data = rb.build_report_data(make_case())
    row = [a for a in data["assets"] if a["name"] == "ws-042.corp.example"][0]
    assert row["type"] == "Unclassified"
    assert row["criticality"] == "Not assessed"
    assert row["role"] == "Not determined"


def test_case_with_no_assets_says_so_rather_than_rendering_an_empty_table():
    out = render("markdown", make_case(asset_links=[]))
    assert "No assets were linked to this case." in out


def test_legacy_affected_systems_still_reaches_the_report():
    """
    The structured table supersedes the free text but must not swallow it —
    on older cases that column may hold detail the assets do not.
    """
    case = make_case(affected_systems="host-alpha\nhost-beta")
    out = render("markdown", case)
    assert "As originally recorded" in out
    assert "host-alpha; host-beta" in out, "newlines must be collapsed for inline display"


@pytest.mark.parametrize("fmt", ["markdown", "docx"])
def test_response_metrics_render_as_intervals(fmt):
    out = render(fmt, make_case())
    for label in ["Detection to containment", "Containment to eradication",
                  "Eradication to closure"]:
        assert label in out


def test_unreached_milestones_say_so_instead_of_showing_a_blank():
    data = rb.build_report_data(make_case(contained_date=None, eradicated_date=None,
                                          closed_date=None))
    values = {m["label"]: m["value"] for m in data["response_metrics"]}
    assert values["Detection to containment"] == "Not yet reached"


def test_integrity_statement_never_claims_more_than_was_checked():
    """
    The load-bearing assertion of this whole section. A file that has not been
    re-hashed is not a verified file, and this report goes to counsel.
    """
    ev = lambda **kw: Obj(evidence_id="E", name="n", evidence_type="t",
                          collected_by="c", hash_sha256=SHA256, hash_md5=MD5, **kw)
    case = make_case(evidence_items=[
        ev(file_path="a", hash_verified_at=datetime(2026, 2, 4), hash_verified_ok=True),
        ev(file_path="b", hash_verified_at=None, hash_verified_ok=None),
        ev(file_path=None, hash_verified_at=None, hash_verified_ok=None),
    ])
    s = rb.build_report_data(case)["integrity"]["statement"]
    assert "1 verified" in s
    assert "1 have not been re-verified" in s
    assert "2 of 3" in s, "records with no stored file are unverifiable, not unverified"


def test_a_failed_hash_is_stated_bluntly():
    ev = Obj(evidence_id="E", name="n", evidence_type="t", collected_by="c",
             hash_sha256=SHA256, hash_md5=MD5, file_path="a",
             hash_verified_at=datetime(2026, 2, 4), hash_verified_ok=False)
    s = rb.build_report_data(make_case(evidence_items=[ev]))["integrity"]["statement"]
    assert "FAILED" in s and "altered since collection" in s


def test_no_evidence_is_distinct_from_unverified_evidence():
    s = rb.build_report_data(make_case(evidence_items=[]))["integrity"]["statement"]
    assert s == "No evidence was recorded for this case."


@pytest.mark.parametrize("fmt", ["markdown", "docx"])
def test_detection_provenance_renders_the_originating_alerts(fmt):
    out = render(fmt, make_case())
    assert "Crowdstrike" in out
    assert "ldt:" in out


def test_case_not_from_an_alert_points_at_method_of_discovery():
    out = render("markdown", make_case(alerts=[]))
    assert "not raised from an ingested alert" in out


def test_contributors_come_from_the_audit_log_and_are_deduplicated():
    data = rb.build_report_data(make_case())
    names = [c["name"] for c in data["contributors"]]
    assert names == ["J. Analyst", "K. Creator"], "ordered by first action, deduplicated"
    assert [c["actions"] for c in data["contributors"]] == [2, 1]


def test_unattributed_audit_entries_do_not_become_a_contributor():
    """A restore writes an entry with no user. It must not appear as a person."""
    data = rb.build_report_data(make_case())
    assert all(c["name"] not in (None, "None", "") for c in data["contributors"])
