"""
Incident Report (AAR) generation for a single case.

Builds one plain-data dict from a Case and its related rows, then renders
that same dict to Markdown and to a Word document. Both renderers consume
`build_report_data()` output exclusively — neither one reaches back into
the ORM — so the two formats can never quietly drift apart from each other.

Structure and required content are traced to the Incident Management Plan,
Phase VI (Lessons Learned): chronological event sequence, method of
discovery, preventive measures implemented, and assessment of recovery
sufficiency are named there as mandatory. The severity classification
matrix is IMP Phase II. Closure sign-off authority is IMP §4.2 (CISO) and
§4.4 (IR Commander).
"""

import io
from datetime import date, datetime, timedelta

from app.models import utcnow

# ---------------------------------------------------------------------------
# IMP Phase II severity classification matrix
# ---------------------------------------------------------------------------

# (functional_impact, informational_impact) -> IMP severity label.
# None/None ("no effect, nothing accessed") isn't a codified incident under
# the matrix and has no cell — it stays unmapped on purpose.
_IMP_SEVERITY_MATRIX = {
    ("None", "Limited"): "Sev. 3", ("None", "Moderate"): "Sev. 2", ("None", "Critical"): "Sev. 1",
    ("Limited", "None"): "Sev. 3", ("Limited", "Limited"): "Sev. 3", ("Limited", "Moderate"): "Sev. 2", ("Limited", "Critical"): "Sev. 1",
    ("Moderate", "None"): "Sev. 2", ("Moderate", "Limited"): "Sev. 2", ("Moderate", "Moderate"): "Sev. 2", ("Moderate", "Critical"): "Sev. 1",
    ("Critical", "None"): "Sev. 1", ("Critical", "Limited"): "Sev. 1", ("Critical", "Moderate"): "Sev. 1", ("Critical", "Critical"): "Sev. 1",
}

# IMP §4.2 (CISO) / §4.4 (IR Commander): closure approval authority by severity.
_APPROVAL_AUTHORITY = {
    "Sev. 1": "CISO",
    "Sev. 2": "CISO or IR Commander",
    "Sev. 3": "IR Commander",
}


def imp_severity(functional_impact, informational_impact):
    """Return 'Sev. 1' / 'Sev. 2' / 'Sev. 3', or None if either axis is unset or unmapped."""
    if not functional_impact or not informational_impact:
        return None
    return _IMP_SEVERITY_MATRIX.get((functional_impact, informational_impact))


def approval_authority(imp_sev):
    return _APPROVAL_AUTHORITY.get(imp_sev)


# ---------------------------------------------------------------------------
# Small formatting helpers
# ---------------------------------------------------------------------------

def _fmt_dt(dt, fmt="%Y-%m-%d %H:%M"):
    return dt.strftime(fmt) if dt else None


def _fmt_date(d, fmt="%Y-%m-%d"):
    return d.strftime(fmt) if d else None


def _duration_str(start, end):
    if not start or not end or end < start:
        return None
    delta = end - start
    days, seconds = delta.days, delta.seconds
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    parts = []
    if days:
        parts.append(f"{days} day{'s' if days != 1 else ''}")
    if hours:
        parts.append(f"{hours}h")
    if not days and minutes:
        parts.append(f"{minutes}m")
    return ", ".join(parts) if parts else "< 1m"


def _business_days_between(start, end):
    """Count business days (Mon-Fri) strictly between two dates, inclusive of end."""
    if not start or not end or end < start:
        return None
    count = 0
    d = start
    while d < end:
        d += timedelta(days=1)
        if d.weekday() < 5:
            count += 1
    return count


def _or_placeholder(value, placeholder="Not yet documented."):
    if value is None:
        return placeholder
    if isinstance(value, str) and not value.strip():
        return placeholder
    return value


# XML 1.0 permits only #x9, #xA, #xD and #x20+ in this range. python-docx
# raises ValueError on anything else rather than sanitizing, so a single stray
# control character anywhere in the case text failed the whole .docx download
# with a 500 while the Markdown export succeeded — the analyst saw a broken
# button, not a bad character. \x0b in particular is Word's own soft line
# break, so text pasted out of Word, a PDF, or an alert body carries it in
# routinely.
_XML_UNSAFE_TABLE = str.maketrans({c: None for c in range(32) if c not in (9, 10)})


def _integrity_statement(i: dict) -> str:
    """
    One sentence about evidence integrity that an examiner can rely on.

    It must never claim more than the verification actually covered. A file that
    has never been re-hashed is not a verified file, and a record with nothing
    stored in CAIRN is not unverified — it is unverifiable by this system, which
    is a different statement and the honest one to make.
    """
    if not i["records"]:
        return "No evidence was recorded for this case."

    noun = "record" if i["records"] == 1 else "records"
    if not i["with_files"]:
        return (
            f"{i['records']} evidence {noun}, none with a file stored in CAIRN. "
            f"Their integrity cannot be attested by this system."
        )

    parts = [f"{i['with_files']} of {i['records']} evidence {noun} have a file stored in CAIRN."]
    if i["mismatched"]:
        parts.append(
            f"{i['mismatched']} FAILED hash verification and must be treated as "
            f"altered since collection."
        )
    if i["verified"]:
        parts.append(
            f"{i['verified']} verified against the hash recorded at collection."
        )
    if i["unchecked"]:
        parts.append(f"{i['unchecked']} have not been re-verified since collection.")
    if i["last_checked"]:
        parts.append(f"Most recent verification: {i['last_checked']} UTC.")
    return " ".join(parts)


def _clean(value):
    r"""Return *value* as a string safe to write into a Word document.

    Drops XML-illegal control characters and collapses CRLF to LF. The CRLF
    step earns its place on its own: browsers submit textarea content with
    CRLF, and python-docx turns \r and \n each into their own <w:br/>, which
    double-spaced every newline the analyst typed.
    """
    if value is None:
        return ""
    text = str(value)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return text.translate(_XML_UNSAFE_TABLE)


def _md_cell(value):
    """Escape *value* for use inside a Markdown table cell.

    An unescaped pipe closes the column early and a newline ends the row, so
    one multi-line description silently tore the rest of the table apart.
    Markdown table cells are single-line by construction; the break becomes
    an inline <br>.
    """
    return _clean(value).replace("|", "\\|").replace("\n", "<br>")


# ---------------------------------------------------------------------------
# Data assembly — the single source of truth both renderers read from
# ---------------------------------------------------------------------------

def build_report_data(case):
    """Assemble the full report content for *case* as a plain dict."""

    imp_sev = imp_severity(case.imp_functional_impact, case.imp_informational_impact)
    authority = approval_authority(imp_sev)

    # `.all()` if this is a lazy="dynamic" relationship (AppenderQuery), plain
    # `list()` if the caller passed an ordinary iterable (e.g. a test double)
    # — either way the explicit sort below is what actually orders it.
    timeline_events = list(case.timeline_events.all() if hasattr(case.timeline_events, "all") else case.timeline_events)
    iocs = list(case.iocs.all() if hasattr(case.iocs, "all") else case.iocs)
    evidence_items = list(case.evidence_items)
    status_history = list(case.status_history)
    deviations = list(case.deviations)
    recommendations = list(case.recommendations)

    # Chronological, oldest first, regardless of how the caller passed events in.
    timeline_events = sorted(timeline_events, key=lambda e: e.event_datetime)
    status_history = sorted(status_history, key=lambda s: s.recorded_at or utcnow())
    recommendations = sorted(recommendations, key=lambda r: r.id)

    timeline = []
    for ev in timeline_events:
        # Tactic, technique ID and the readable technique name. The name used
        # to be dropped, which left the reader holding a bare "T1566.002" and
        # a lookup to go do.
        mitre_parts = [x for x in (ev.mitre_tactic, ev.mitre_technique_id, ev.mitre_technique) if x]
        mitre = " — ".join(mitre_parts) if mitre_parts else "—"
        timeline.append({
            "time": _fmt_dt(ev.event_datetime),
            "phase": ev.category or "—",
            "description": ev.description,
            "mitre": mitre,
        })

    ioc_rows = [{
        "type": i.ioc_type,
        "value": i.value,
        "description": i.description or "—",
        "confidence": i.confidence or "—",
    } for i in iocs]

    # Hashes are carried in full, not as a prefix. A 12-character stub reads
    # fine on screen but cannot be verified against anything, and this report
    # leaves the building — to counsel, to an insurer, to an examiner. They
    # ride in their own block below the table rather than as a sixth column,
    # because 64 characters will not sit in a portrait-page cell.
    evidence_rows = [{
        "evidence_id": e.evidence_id,
        "name": e.name,
        "type": e.evidence_type or "—",
        "collected_by": e.collected_by or "—",
        "sha256": e.hash_sha256 or "",
        "md5": e.hash_md5 or "",
    } for e in evidence_items]

    evidence_hashes = [{
        "evidence_id": e["evidence_id"],
        "algorithm": "SHA-256" if e["sha256"] else "MD5",
        "digest": e["sha256"] or e["md5"],
    } for e in evidence_rows if e["sha256"] or e["md5"]]

    # ── Affected assets ──────────────────────────────────────────────────────
    # Supersedes the affected_systems free-text line, which is still printed
    # below it while both are live. A typed, deduplicated list answers what the
    # text never could: what kind of estate was hit, and what each host was in
    # this particular incident.
    asset_rows = [{
        "name": link.asset.name,
        "type": link.asset.type_label,
        "criticality": link.asset.criticality or "Not assessed",
        "role": link.role or "Not determined",
        "notes": _clean(link.notes) or "",
    } for link in sorted(case.asset_links, key=lambda l: l.asset.name.lower())]

    # Computed from the table above, never typed, so the two cannot disagree.
    asset_scope = None
    if asset_rows:
        by_type = {}
        for r in asset_rows:
            by_type[r["type"]] = by_type.get(r["type"], 0) + 1
        breakdown = ", ".join(
            f"{t} ({n})"
            for t, n in sorted(by_type.items(), key=lambda kv: (-kv[1], kv[0]))
        )
        n = len(asset_rows)
        asset_scope = f"{n} asset{'s' if n != 1 else ''} — {breakdown}"
        critical = sum(1 for r in asset_rows if r["criticality"] == "Critical")
        if critical:
            asset_scope += f"; {critical} rated Critical"
        unclassified = by_type.get("Unclassified", 0)
        if unclassified:
            # Said out loud rather than left as a blank cell. An unclassified
            # asset in a report going to leadership is an open question, not a
            # formatting artefact.
            asset_scope += (
                f". {unclassified} not yet classified — the type is unknown, "
                f"not absent"
            )

    # ── Response metrics ─────────────────────────────────────────────────────
    # All four timestamps were already stored and only the total was printed.
    # The intervals are what a reader actually asks for. Derived, not entered,
    # so they cannot contradict the dates they come from.
    response_metrics = [
        {"label": label, "value": value or "Not yet reached"}
        for label, value in (
            ("Detection to containment", _duration_str(case.opened_date, case.contained_date)),
            ("Containment to eradication", _duration_str(case.contained_date, case.eradicated_date)),
            ("Eradication to closure", _duration_str(case.eradicated_date, case.closed_date)),
            ("Total — opened to closed", _duration_str(case.opened_date, case.closed_date)),
        )
    ]

    # ── Evidence integrity ───────────────────────────────────────────────────
    # Listing hashes without saying whether they still match is the weaker half
    # of the chain-of-custody design. This report leaves the building.
    stored_evidence = [e for e in evidence_items if e.file_path]
    last_checked = max(
        (e.hash_verified_at for e in stored_evidence if e.hash_verified_at),
        default=None,
    )
    integrity = {
        "records": len(evidence_items),
        "with_files": len(stored_evidence),
        "verified": sum(1 for e in stored_evidence if e.hash_verified_ok is True),
        "mismatched": sum(1 for e in stored_evidence if e.hash_verified_ok is False),
        "unchecked": sum(1 for e in stored_evidence if e.hash_verified_ok is None),
        "last_checked": _fmt_dt(last_checked),
    }
    integrity["statement"] = _integrity_statement(integrity)

    # ── Detection provenance ─────────────────────────────────────────────────
    # method_of_discovery is what an analyst wrote afterwards. This is what the
    # tooling actually did. The two disagreeing is itself worth seeing.
    alerts_sorted = sorted(
        case.alerts,
        key=lambda a: (a.cs_created_at or a.fetched_at or datetime.max),
    )
    alert_rows = [{
        "source": (a.source or "").replace("_", " ").title() or "—",
        "external_id": a.external_id or "—",
        "severity": a.severity_name or "—",
        "host": a.host_hostname or a.host_ip or "—",
        "detected": _fmt_dt(a.cs_created_at) or "—",
        "ingested": _fmt_dt(a.fetched_at) or "—",
        "promoted_by": a.reviewed_by.name if a.reviewed_by else "—",
    } for a in alerts_sorted]

    # ── Contributors ─────────────────────────────────────────────────────────
    # Everyone who touched the case, from the audit log. The lead analyst is
    # named elsewhere; this is who actually did the work, which is not always
    # the same person and is standard AAR content.
    seen = {}
    for entry in case.audit_entries:
        if not entry.changed_by_id:
            continue
        name = entry.changed_by.name if entry.changed_by else f"user {entry.changed_by_id}"
        row = seen.setdefault(name, {"name": name, "actions": 0, "first": None, "last": None})
        row["actions"] += 1
        when = entry.changed_at
        if when:
            if row["first"] is None or when < row["first"]:
                row["first"] = when
            if row["last"] is None or when > row["last"]:
                row["last"] = when
    contributor_rows = [{
        "name": r["name"],
        "actions": r["actions"],
        "first": _fmt_dt(r["first"]) or "—",
        "last": _fmt_dt(r["last"]) or "—",
    } for r in sorted(seen.values(), key=lambda r: (r["first"] or datetime.max))]

    deviation_rows = [{
        "deviation": d.deviation,
        "standard_procedure": d.standard_procedure or "—",
        "justification": d.justification or "—",
        "approved_by": d.approved_by or "—",
    } for d in deviations]

    recommendation_rows = [{
        "n": idx + 1,
        "text": r.text,
        "disposition": r.disposition,
        "owner": r.owner or "—",
        "target": _fmt_date(r.target_date) or "—",
        "rtp_ref": r.risk_treatment_ref or "Pending",
        "status": r.status,
        "justification": (r.risk_acceptance_justification or "").strip(),
    } for idx, r in enumerate(recommendations)]

    # IMP Phase VI requires a risk acceptance to carry a documented
    # justification, and the Report tab refuses to save one without it. The
    # export dropped it anyway, so the AAR showed "Risk Acceptance" with the
    # reasoning nowhere in the document. Pulled out here so every renderer
    # states it beneath the table without widening the table to hold it.
    risk_acceptances = [
        {"n": r["n"], "justification": r["justification"]}
        for r in recommendation_rows
        if r["disposition"] == "Risk Acceptance" and r["justification"]
    ]

    status_history_rows = [{
        "date": _fmt_dt(s.recorded_at),
        "change": f"{s.old_status or '—'} → {s.new_status}",
        "notes": s.notes or "—",
    } for s in status_history]

    ll_sla_met = None
    ll_business_days = None
    if case.lessons_learned_date and case.closed_date:
        ll_business_days = _business_days_between(case.closed_date.date(), case.lessons_learned_date)
        if ll_business_days is not None:
            ll_sla_met = ll_business_days <= 5

    return {
        "case_id": case.case_id,
        "title": case.title,
        "case_type": case.case_type or "—",
        "cairn_severity": case.severity,
        "imp_functional_impact": case.imp_functional_impact,
        "imp_informational_impact": case.imp_informational_impact,
        "imp_severity": imp_sev,
        "approval_authority": authority,
        "status": case.status,
        "lead_analyst": case.lead_analyst.name if case.lead_analyst else "Unassigned",
        "escalated": case.escalated,
        "board_flagged": case.board_flagged,
        "opened": _fmt_dt(case.opened_date),
        "contained": _fmt_dt(case.contained_date),
        "eradicated": _fmt_dt(case.eradicated_date),
        "closed": _fmt_dt(case.closed_date),
        "duration": _duration_str(case.opened_date, case.closed_date),

        "description": _or_placeholder(case.description, "No executive summary recorded."),
        "method_of_discovery": _or_placeholder(case.method_of_discovery),
        "initial_vector": _or_placeholder(case.initial_vector),
        "affected_systems": _or_placeholder(case.affected_systems, "Not recorded."),
        # The raw column, collapsed to a single line. Renderers show it only when
        # it has content, under the structured asset table rather than instead of
        # it. Newlines are joined rather than kept: this renders inline after a
        # bold label in all three outputs, and a line break there breaks out of
        # the label in Markdown and out of the paragraph in Word.
        "affected_systems_text": "; ".join(
            line.strip()
            for line in (case.affected_systems or "").splitlines()
            if line.strip()
        ),
        "affected_users": _or_placeholder(case.affected_users, "Not recorded."),
        "estimated_impact": _or_placeholder(case.estimated_impact, "Not recorded."),

        "timeline": timeline,

        "iocs": ioc_rows,
        "assets": asset_rows,
        "asset_scope": asset_scope,
        "response_metrics": response_metrics,
        "integrity": integrity,
        "alerts": alert_rows,
        "contributors": contributor_rows,
        "evidence": evidence_rows,
        "evidence_hashes": evidence_hashes,

        "root_cause": _or_placeholder(case.root_cause),

        "recovery_assessment": _or_placeholder(case.recovery_assessment),
        "recovery_sufficient": case.recovery_sufficient,

        "deviations": deviation_rows,

        "lessons_learned_date": _fmt_date(case.lessons_learned_date),
        "lessons_learned_attendees": _or_placeholder(case.lessons_learned_attendees, "Not recorded."),
        "lessons_learned_notes": _or_placeholder(case.lessons_learned_notes),
        "lessons_learned_sla_met": ll_sla_met,
        "lessons_learned_business_days": ll_business_days,

        "recommendations": recommendation_rows,
        "risk_acceptances": risk_acceptances,

        "status_history": status_history_rows,

        "prepared_by": case.lead_analyst.name if case.lead_analyst else (case.created_by.name if case.created_by else "—"),
        "generated_at": utcnow().strftime("%Y-%m-%d %H:%M UTC"),
    }


# ---------------------------------------------------------------------------
# Markdown renderer
# ---------------------------------------------------------------------------

def render_markdown(data: dict) -> str:
    lines = []
    a = lines.append

    a(f"# After Action Report — {data['case_id']}")
    a("")
    a(f"## {data['title']}")
    a("")
    a("| | |")
    a("|---|---|")
    a(f"| **Case ID** | {data['case_id']} |")
    a(f"| **Classification** | {data['case_type']} |")
    a(f"| **Cairn Severity** | {data['cairn_severity']} |")
    if data["imp_severity"]:
        detail = f"{data['imp_severity']} (Functional Impact: {data['imp_functional_impact']}; Informational Impact: {data['imp_informational_impact']})"
    else:
        detail = "Not yet classified — set Functional and Informational Impact on the Report tab."
    a(f"| **IMP Classification** | {detail} |")
    a(f"| **Status** | {data['status']} |")
    a(f"| **Lead Analyst** | {data['lead_analyst']} |")
    esc = "Yes" if data["escalated"] else "No"
    if data["board_flagged"]:
        esc += " — Board notified"
    a(f"| **Escalated** | {esc} |")
    a(f"| **Opened** | {data['opened'] or '—'} |")
    a(f"| **Contained** | {data['contained'] or '—'} |")
    a(f"| **Eradicated** | {data['eradicated'] or '—'} |")
    a(f"| **Closed** | {data['closed'] or '—'} |")
    a(f"| **Total Duration** | {data['duration'] or '—'} |")
    a("")
    a("---")
    a("")

    a("## Executive Summary")
    a("")
    a(data["description"])
    a("")

    a("## Incident Overview")
    a("")
    a(f"**Method of Discovery:** {data['method_of_discovery']}")
    a("")
    a(f"**Initial Vector:** {data['initial_vector']}")
    a("")
    a(f"**Affected Users:** {data['affected_users']}")
    a("")
    a(f"**Estimated Impact:** {data['estimated_impact']}")
    a("")

    a("## Affected Assets")
    a("")
    if data["assets"]:
        a(f"**Scope:** {data['asset_scope']}")
        a("")
        a("| Asset | Type | Criticality | Role in this incident | Notes |")
        a("|---|---|---|---|---|")
        for s in data["assets"]:
            a(f"| {_md_cell(s['name'])} | {_md_cell(s['type'])} | "
              f"{_md_cell(s['criticality'])} | {_md_cell(s['role'])} | "
              f"{_md_cell(s['notes'])} |")
    else:
        a("No assets were linked to this case.")
    a("")
    if data["affected_systems_text"]:
        # Kept because it is what an analyst actually typed, and on older cases
        # it may hold detail the structured list does not.
        a(f"**As originally recorded:** {data['affected_systems_text']}")
        a("")

    a("## Response Metrics")
    a("")
    a("| Interval | Elapsed |")
    a("|---|---|")
    for m in data["response_metrics"]:
        a(f"| {_md_cell(m['label'])} | {_md_cell(m['value'])} |")
    a("")

    a("## Detection Provenance")
    a("")
    if data["alerts"]:
        a("| Source | Alert ID | Severity | Host | Detected (UTC) | Ingested (UTC) | Promoted by |")
        a("|---|---|---|---|---|---|---|")
        for al in data["alerts"]:
            a(f"| {_md_cell(al['source'])} | {_md_cell(al['external_id'])} | "
              f"{_md_cell(al['severity'])} | {_md_cell(al['host'])} | "
              f"{_md_cell(al['detected'])} | {_md_cell(al['ingested'])} | "
              f"{_md_cell(al['promoted_by'])} |")
    else:
        a("This case was not raised from an ingested alert. See Method of "
          "Discovery above for how it came to attention.")
    a("")

    a("## Timeline of Events")
    a("")
    if data["timeline"]:
        a("| Time (UTC) | Phase | Description | MITRE ATT&CK |")
        a("|---|---|---|---|")
        for t in data["timeline"]:
            a(f"| {_md_cell(t['time'])} | {_md_cell(t['phase'])} | "
              f"{_md_cell(t['description'])} | {_md_cell(t['mitre'])} |")
    else:
        a("No timeline events recorded.")
    a("")

    a("## Indicators of Compromise")
    a("")
    if data["iocs"]:
        a("| Type | Value | Description | Confidence |")
        a("|---|---|---|---|")
        for i in data["iocs"]:
            a(f"| {_md_cell(i['type'])} | {_md_cell(i['value'])} | "
              f"{_md_cell(i['description'])} | {_md_cell(i['confidence'])} |")
    else:
        a("No IOCs recorded.")
    a("")

    a("## Evidence Collected")
    a("")
    if data["evidence"]:
        a("| Evidence ID | Name | Type | Collected By |")
        a("|---|---|---|---|")
        for e in data["evidence"]:
            a(f"| {_md_cell(e['evidence_id'])} | {_md_cell(e['name'])} | "
              f"{_md_cell(e['type'])} | {_md_cell(e['collected_by'])} |")
        a("")
        if data["evidence_hashes"]:
            a("**Integrity hashes**")
            a("")
            for h in data["evidence_hashes"]:
                a(f"- `{h['evidence_id']}` — {h['algorithm']}: `{h['digest']}`")
            a("")
        a("*Full chain-of-custody detail is retained in Cairn per each evidence record.*")
    else:
        a("No evidence recorded.")
    a("")

    a("### Evidence Integrity")
    a("")
    a(data["integrity"]["statement"])
    a("")

    a("## Who Worked This Incident")
    a("")
    a(f"**Lead analyst:** {data['lead_analyst']}")
    a("")
    if data["contributors"]:
        a("| Contributor | Recorded actions | First (UTC) | Last (UTC) |")
        a("|---|---|---|---|")
        for cr in data["contributors"]:
            a(f"| {_md_cell(cr['name'])} | {cr['actions']} | "
              f"{_md_cell(cr['first'])} | {_md_cell(cr['last'])} |")
        a("")
        a("Derived from the audit log — actions recorded in CAIRN only. Work done "
          "outside the console is not represented here.")
    else:
        a("No attributed actions recorded against this case.")
    a("")

    a("## Root Cause")
    a("")
    a(data["root_cause"])
    a("")

    a("## Recovery Sufficiency Assessment")
    a("")
    a(f"**Assessment: {data['recovery_sufficient'] or 'Not yet assessed'}.**")
    a("")
    a(data["recovery_assessment"])
    a("")

    a("## Deviations from Standard Procedure")
    a("")
    if data["deviations"]:
        a("| Deviation | Standard Procedure | Justification | Approved By |")
        a("|---|---|---|---|")
        for d in data["deviations"]:
            a(f"| {_md_cell(d['deviation'])} | {_md_cell(d['standard_procedure'])} | "
              f"{_md_cell(d['justification'])} | {_md_cell(d['approved_by'])} |")
    else:
        a("No deviations from standard procedure recorded.")
    a("")

    a("## Lessons Learned Meeting")
    a("")
    if data["lessons_learned_date"]:
        sla = ""
        if data["lessons_learned_sla_met"] is True:
            sla = f" ({data['lessons_learned_business_days']} business days post-closure — within the IMP's 5-business-day requirement)"
        elif data["lessons_learned_sla_met"] is False:
            sla = f" ({data['lessons_learned_business_days']} business days post-closure — outside the IMP's 5-business-day requirement)"
        a(f"Held {data['lessons_learned_date']}{sla}.")
        a("")
        a(f"**Attendees:** {data['lessons_learned_attendees']}")
    else:
        a("Not yet held or not yet recorded. IMP Phase VI requires this meeting within 5 business days of closure.")
    a("")
    a(data["lessons_learned_notes"])
    a("")

    a("## Recommendations")
    a("")
    if data["recommendations"]:
        a("| # | Recommendation | Disposition | Owner | Target | RTP Ref. | Status |")
        a("|---|---|---|---|---|---|---|")
        for r in data["recommendations"]:
            a(f"| {r['n']} | {_md_cell(r['text'])} | {_md_cell(r['disposition'])} | "
              f"{_md_cell(r['owner'])} | {_md_cell(r['target'])} | "
              f"{_md_cell(r['rtp_ref'])} | {_md_cell(r['status'])} |")
        a("")
        if data["risk_acceptances"]:
            a("**Documented risk acceptances (IMP Phase VI)**")
            a("")
            for ra in data["risk_acceptances"]:
                a(f"- **#{ra['n']}:** {ra['justification']}")
            a("")
        a("*Per IMP Phase VI, every identified gap must resolve to remediation, a compensating control, or a formal, documented risk acceptance, and be logged against the organizational risk treatment plan.*")
    else:
        a("No recommendations recorded.")
    a("")

    a("## Case Status History")
    a("")
    if data["status_history"]:
        a("| Date/Time | Change | Notes |")
        a("|---|---|---|")
        for s in data["status_history"]:
            a(f"| {_md_cell(s['date'])} | {_md_cell(s['change'])} | {_md_cell(s['notes'])} |")
    else:
        a("No status changes recorded.")
    a("")
    a("---")
    a("")

    a("## Approval")
    a("")
    if data["imp_severity"]:
        a(f"**Approval authority (per IMP §4.2 / §4.4):** This incident is classified {data['imp_severity']}. Closure requires approval from: {data['approval_authority']}.")
    else:
        a("**Approval authority:** Not yet determined — set Functional and Informational Impact on the Report tab to compute the IMP severity classification and required approver.")
    a("")
    a(f"*Report generated from Cairn case {data['case_id']} on {data['generated_at']}. Prepared by {data['prepared_by']}.*")
    a("")
    a("Approved by: ____________________________&nbsp;&nbsp;&nbsp;&nbsp;Date: ______________")
    a("")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# DOCX renderer
# ---------------------------------------------------------------------------

def render_docx(data: dict) -> io.BytesIO:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY = RGBColor(0x1F, 0x38, 0x64)
    RED = RGBColor(0xC0, 0x00, 0x00)
    GREY = RGBColor(0x59, 0x59, 0x59)
    LIGHTGREY = "F2F2F2"
    AMBER = "FDEBD0"

    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(0.75)
    section.top_margin = section.bottom_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(20)
    h1.font.color.rgb = NAVY
    h1.font.bold = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(14)
    h2.font.color.rgb = NAVY
    h2.font.bold = True

    def shade_cell(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def set_cell_text(cell, text, bold=False, color=None, size=9.5, italic=False):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(_clean(text))
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color

    def add_table(headers, rows, widths=None):
        """Render a table. *widths* is an optional list of Inches per column.

        Word's autofit divides the page evenly and then squeezes, which turns
        a long narrative column into a ribbon next to five short ones. Passing
        explicit widths gives the prose the room and holds the label columns
        to what they actually need.
        """
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = widths is None
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            set_cell_text(hdr[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=9.5)
            shade_cell(hdr[i], "1F3864")
        for ri, row in enumerate(rows):
            cells = table.add_row().cells
            for i, val in enumerate(row):
                set_cell_text(cells[i], val, size=9.5)
                if ri % 2 == 1:
                    shade_cell(cells[i], LIGHTGREY)
        if widths:
            # Width has to be set on every cell in a column, not just the
            # column object — Word reads the cell-level value.
            for row in table.rows:
                for i, w in enumerate(widths[:len(headers)]):
                    row.cells[i].width = w
        doc.add_paragraph()
        return table

    def add_info_table(pairs):
        table = doc.add_table(rows=0, cols=2)
        table.autofit = True
        for label, value, *shade in pairs:
            row = table.add_row().cells
            set_cell_text(row[0], label, bold=True, size=9.5)
            shade_cell(row[0], LIGHTGREY)
            set_cell_text(row[1], value, size=9.5)
            if shade:
                shade_cell(row[1], shade[0])
        doc.add_paragraph()

    def add_heading(text):
        doc.add_heading(text, level=2)

    def add_para(text, bold=False, italic=False, color=None, size=None):
        p = doc.add_paragraph()
        run = p.add_run(_clean(text))
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
        if size:
            run.font.size = Pt(size)
        return p

    # ── Title block ──────────────────────────────────────────────────────
    kicker = doc.add_paragraph()
    kr = kicker.add_run("AFTER ACTION REPORT")
    kr.bold = True
    kr.font.color.rgb = RED
    kr.font.size = Pt(10)

    doc.add_heading(f"{data['case_id']} — {data['title']}", level=1)

    imp_detail = (
        f"{data['imp_severity']} (Functional Impact: {data['imp_functional_impact']}; "
        f"Informational Impact: {data['imp_informational_impact']})"
        if data["imp_severity"]
        else "Not yet classified — set Functional and Informational Impact on the Report tab."
    )
    esc = "Yes" if data["escalated"] else "No"
    if data["board_flagged"]:
        esc += " — Board notified"

    add_info_table([
        ("Case ID", data["case_id"]),
        ("Classification", data["case_type"]),
        ("Cairn Severity", data["cairn_severity"]),
        ("IMP Classification", imp_detail, AMBER),
        ("Status", data["status"]),
        ("Lead Analyst", data["lead_analyst"]),
        ("Escalated", esc),
        ("Opened", data["opened"] or "—"),
        ("Contained", data["contained"] or "—"),
        ("Eradicated", data["eradicated"] or "—"),
        ("Closed", data["closed"] or "—"),
        ("Total Duration", data["duration"] or "—"),
    ])

    add_heading("Executive Summary")
    add_para(data["description"])

    add_heading("Incident Overview")
    add_para(f"Method of Discovery: {data['method_of_discovery']}")
    add_para(f"Initial Vector: {data['initial_vector']}")
    add_para(f"Affected Users: {data['affected_users']}")
    add_para(f"Estimated Impact: {data['estimated_impact']}")

    add_heading("Affected Assets")
    if data["assets"]:
        add_para(data["asset_scope"], bold=True)
        add_table(
            ["Asset", "Type", "Criticality", "Role in this incident", "Notes"],
            [[s["name"], s["type"], s["criticality"], s["role"], s["notes"]]
             for s in data["assets"]],
            widths=[Inches(1.7), Inches(1.3), Inches(0.9), Inches(1.5), Inches(1.6)],
        )
    else:
        add_para("No assets were linked to this case.", italic=True, color=GREY)
    if data["affected_systems_text"]:
        add_para(f"As originally recorded: {data['affected_systems_text']}",
                 italic=True, color=GREY)

    add_heading("Response Metrics")
    add_table(
        ["Interval", "Elapsed"],
        [[m["label"], m["value"]] for m in data["response_metrics"]],
        widths=[Inches(3.4), Inches(3.6)],
    )

    add_heading("Detection Provenance")
    if data["alerts"]:
        add_table(
            ["Source", "Alert ID", "Severity", "Host", "Detected", "Ingested", "Promoted by"],
            [[al["source"], al["external_id"], al["severity"], al["host"],
              al["detected"], al["ingested"], al["promoted_by"]]
             for al in data["alerts"]],
            widths=[Inches(0.9), Inches(1.5), Inches(0.8), Inches(1.2),
                    Inches(1.0), Inches(1.0), Inches(0.6)],
        )
    else:
        add_para(
            "This case was not raised from an ingested alert. See Method of "
            "Discovery above for how it came to attention.",
            italic=True, color=GREY,
        )

    add_heading("Timeline of Events")
    if data["timeline"]:
        add_table(
            ["Time (UTC)", "Phase", "Description", "MITRE ATT&CK"],
            [[t["time"], t["phase"], t["description"], t["mitre"]] for t in data["timeline"]],
            widths=[Inches(1.1), Inches(1.0), Inches(3.1), Inches(1.8)],
        )
    else:
        add_para("No timeline events recorded.", italic=True, color=GREY)

    add_heading("Indicators of Compromise")
    if data["iocs"]:
        add_table(
            ["Type", "Value", "Description", "Confidence"],
            [[i["type"], i["value"], i["description"], i["confidence"]] for i in data["iocs"]],
            widths=[Inches(1.0), Inches(2.2), Inches(2.9), Inches(0.9)],
        )
    else:
        add_para("No IOCs recorded.", italic=True, color=GREY)

    add_heading("Evidence Collected")
    if data["evidence"]:
        add_table(
            ["Evidence ID", "Name", "Type", "Collected By"],
            [[e["evidence_id"], e["name"], e["type"], e["collected_by"]] for e in data["evidence"]],
            widths=[Inches(1.1), Inches(2.6), Inches(1.6), Inches(1.7)],
        )
        if data["evidence_hashes"]:
            add_para("Integrity hashes", bold=True, size=9.5)
            for h in data["evidence_hashes"]:
                p = doc.add_paragraph()
                lbl = p.add_run(f"{h['evidence_id']} — {h['algorithm']}: ")
                lbl.font.size = Pt(8.5)
                dig = p.add_run(_clean(h["digest"]))
                dig.font.name = "Consolas"
                dig.font.size = Pt(8.5)
        add_para(
            "Full chain-of-custody detail is retained in Cairn per each evidence record.",
            italic=True, color=GREY, size=8.5,
        )
    else:
        add_para("No evidence recorded.", italic=True, color=GREY)

    add_para("Evidence Integrity", bold=True)
    add_para(data["integrity"]["statement"])

    add_heading("Who Worked This Incident")
    add_para(f"Lead analyst: {data['lead_analyst']}")
    if data["contributors"]:
        add_table(
            ["Contributor", "Recorded actions", "First", "Last"],
            [[c["name"], str(c["actions"]), c["first"], c["last"]]
             for c in data["contributors"]],
            widths=[Inches(2.2), Inches(1.2), Inches(1.8), Inches(1.8)],
        )
        add_para(
            "Derived from the audit log — actions recorded in CAIRN only. Work "
            "done outside the console is not represented here.",
            italic=True, color=GREY,
        )
    else:
        add_para("No attributed actions recorded against this case.",
                 italic=True, color=GREY)

    add_heading("Root Cause")
    add_para(data["root_cause"])

    add_heading("Recovery Sufficiency Assessment")
    add_para(f"Assessment: {data['recovery_sufficient'] or 'Not yet assessed'}.", bold=True)
    add_para(data["recovery_assessment"])

    add_heading("Deviations from Standard Procedure")
    if data["deviations"]:
        add_table(
            ["Deviation", "Standard Procedure", "Justification", "Approved By"],
            [[d["deviation"], d["standard_procedure"], d["justification"], d["approved_by"]] for d in data["deviations"]],
            widths=[Inches(1.9), Inches(1.9), Inches(2.3), Inches(0.9)],
        )
    else:
        add_para("No deviations from standard procedure recorded.", italic=True, color=GREY)

    add_heading("Lessons Learned Meeting")
    if data["lessons_learned_date"]:
        sla = ""
        if data["lessons_learned_sla_met"] is True:
            sla = f" ({data['lessons_learned_business_days']} business days post-closure — within the IMP's 5-business-day requirement)"
        elif data["lessons_learned_sla_met"] is False:
            sla = f" ({data['lessons_learned_business_days']} business days post-closure — outside the IMP's 5-business-day requirement)"
        add_para(f"Held {data['lessons_learned_date']}{sla}.")
        add_para(f"Attendees: {data['lessons_learned_attendees']}")
    else:
        add_para(
            "Not yet held or not yet recorded. IMP Phase VI requires this meeting "
            "within 5 business days of closure.",
            italic=True, color=GREY,
        )
    add_para(data["lessons_learned_notes"])

    add_heading("Recommendations")
    if data["recommendations"]:
        add_table(
            ["#", "Recommendation", "Disposition", "Owner", "Target", "RTP Ref.", "Status"],
            [[r["n"], r["text"], r["disposition"], r["owner"], r["target"], r["rtp_ref"], r["status"]]
             for r in data["recommendations"]],
            widths=[Inches(0.35), Inches(2.35), Inches(1.2), Inches(0.9),
                    Inches(0.8), Inches(0.75), Inches(0.65)],
        )
        if data["risk_acceptances"]:
            add_para("Documented risk acceptances (IMP Phase VI)", bold=True, size=9.5)
            for ra in data["risk_acceptances"]:
                add_para(f"#{ra['n']}: {ra['justification']}", size=9)
        add_para(
            "Per IMP Phase VI, every identified gap must resolve to remediation, a compensating "
            "control, or a formal, documented risk acceptance, and be logged against the "
            "organizational risk treatment plan.",
            italic=True, color=GREY, size=8.5,
        )
    else:
        add_para("No recommendations recorded.", italic=True, color=GREY)

    add_heading("Case Status History")
    if data["status_history"]:
        add_table(
            ["Date/Time", "Change", "Notes"],
            [[s["date"], s["change"], s["notes"]] for s in data["status_history"]],
            widths=[Inches(1.3), Inches(1.9), Inches(3.8)],
        )
    else:
        add_para("No status changes recorded.", italic=True, color=GREY)

    add_heading("Approval")
    if data["imp_severity"]:
        add_para(
            f"Approval authority (per IMP §4.2 / §4.4): This incident is classified "
            f"{data['imp_severity']}. Closure requires approval from: {data['approval_authority']}.",
            bold=True,
        )
    else:
        add_para(
            "Approval authority: Not yet determined — set Functional and Informational Impact "
            "on the Report tab to compute the IMP severity classification and required approver.",
            bold=True,
        )
    add_para(
        f"Report generated from Cairn case {data['case_id']} on {data['generated_at']}. "
        f"Prepared by {data['prepared_by']}.",
        italic=True, color=GREY, size=9,
    )
    add_para("Approved by: ____________________________          Date: ______________")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
